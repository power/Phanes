import json
import docx
import docx.document
import docx.text

doc = docx.Document()
doc.add_heading("A Guide to Pwning Phanes", 1) # make a document and make a heading for it

with open("vulns.json") as f:
    vulns = json.load(f) # read the contents of the file to a variable
sids = []
print("> Starting report generation.")
for i in range(20):
    try:
        if (vulns['vulns'][i]['Name'] == "ASREP"): # if the name of the vuln is ASREP
             doc.add_heading("ASREP Roasting", 2)
             if (vulns['vulns'][i]['Status'] == "True"): # if it is present
               for c in range(4):
                         try:
                              #print(vulns['vulns'][i]["SID"+str(c)])
                              sids.append(vulns['vulns'][i]["SID"+str(c)]) # try and get the user SID's and store them ina  list
                         except (KeyError, TypeError): # if something goes wrong, try again
                              continue
               #print(sids)
               asrep_roasting = """
     The process of sending an AS-REQ so that a DC can validate a users authentication, and then reply with an AS-REP is more commonly known as Kerberos pre-authentication and prevents offline password guessing.

     If you take Kerberos pre-authentication out of the picture, attackers could send AS-REQ to the domain on behalf of any AD user. The AS-REP could then be taken by the attacker and used to perform an offline password guessing attack. 

     There is an option in AD environments that by default, is disabled, "Do not require Kerberos preauthentication is disabled". Although, this can be enabled manually and is commonly seen in assessments to be enabled for the account or service to work. We can find accounts on an AD network that are susceptible to AS-REP roasting with the use of `impacket-GetNPUsers`.

     We'll pass `-dc-ip` for the DC's IP, `-request` to signify we are wanting to request the TGT, `-outputfile` for where to output the AS-REP hash. We'll then need to provide a `domain/user` to log in with, which you will have found. 
     """

               asrep_roasting += f"\n\n You should receive the following users password hashes: {sids[0]}, {sids[1]}, {sids[2]}" #add the sids in
               doc.add_paragraph(asrep_roasting) # add all the info
               
               print("> ASREP Roasting Added")
               #print(vulns['vulns'][i]['SID1'])
             elif (vulns['vulns'][i]['Status'] == 'False'): # if as-rep roasting isnt present
                    asrep_roasting = "AS-REP Roasting was not present on this network."
                    doc.add_paragraph(asrep_roasting)
                    print("> ASREP Roasting Added.")
        elif (vulns['vulns'][i]['Name'] == "unconstrainedDelegation"):
             doc.add_heading("Unconstrained Delegation", 2)
             if(vulns['vulns'][i]['Status'] == "False"):
                  doc.add_paragraph("Unconstrained Delegation was not present on the network.")
             else:
               doc.add_paragraph(f"Uncond Status: {vulns['vulns'][i]['Status']}")
               doc.add_paragraph(f"User needed: {vulns['vulns'][i]['SID1']}")
             print("> Unconstrained Delegation Added")
        elif (vulns['vulns'][i]['Name'] == "dcSync"):
          doc.add_heading("DCSync", 2)
          if(vulns['vulns'][i]['SID1'] == "None"):
             doc.add_paragraph("A DCSync attack was not possible on this network.")
          else:
             dcsync = r"""
In live environments, there is usually multiple DC's to provide greater stability across an AD network. Multiple DC's existing in a network will need to remain synchronised and therefore utilise the Directory Replication Service (DRS) Remote Protocol. This protocol utilises replication to synchronise DC's that may be out of sync with the others. DC's may request updates on objects, whether that be services, groups or users through `IDL_DRSGetNCChanges` API call. 

When this request is made by a DC, there is no checks as the DC's assume that no other computer besides DC's would make such a request, luckily for us, this means we can take advantage of it. The only check that other DC's make is that the SID has the appropriate privileges to make such a request, which is something we can still take advantage of. For this kind of attack to work, a user needs to have the following permissions:
- Replicating Directory Changes
- Replicating Directory Changes All
- Replicating Directory Changes in Filtered Set

Unless explicitly specified, members of the following groups will automatically have those permissions:
- Domain Admins
- Enterprise Admins
- Administrators

This means that if we wanted to take advantage of this, we need to compromise an account that has either A. the permission specifically assigned or B. is in one of the specified groups. If we can perform this kind of attack, we're impersonating a DC and can therefore call whatever credentials we want. 

We'll need two devices to perform this attack so will use `mimikatz` on a computer that is on the domain and `impacket-secretsdump` on our own Kali install. 

With mimikatz we'll load it up and execute `lsadump::dcsync /user:fakecompany.local\USER` where `/user:fakecompany.local\USER` can represent any user we choose. From this command, we'll get the NTLM hash which we could pass the `hashcat` but for this attack, we're not wanting to perform lateral movement. 

Assuming we're performing this attack from our Kali machine, we'll pass `-just-dc-user $USER` and then provide `DOMAIN/USER:'PASSWORD'@DC_IP` """
             dcsync += f"\n\nYou should uncover the user {vulns['vulns'][i]['SID1']} from this attack."
             doc.add_paragraph(dcsync)
             print("> DCSync Added")
        elif (vulns['vulns'][i]['Name'] == "Kerberoasting"):
             doc.add_heading("Kerberoasting", 2)
             if(vulns['vulns'][i]['SID1'] == "None"):
                 doc.add_paragraph("Kerberoasting was not possible on the network.")
             else:
               kerberoasting = r"""
     As we've already covered, when a user wants to access a resource that is hosted by a Service Principal Name (SPN), the client requests a service ticket that is generated by the DC. The service ticket is decrypted and validated by the application server since it's encrypted via the password hash of the SPN. When the service ticket is requested from the DC, there are no checks to confirm whether the user has the permissions to access the service hosted by the SPN. This set of steps is performed when the user is attempting to connect to the service it self which means that we know the SPN we want to target and can therefore request a service ticket from the DC. 

     Since the service ticket is encrypted using the SPN's password hash, if we can decrypt the password through bruteforce or guessing we can crack the cleartext password of the service account. This is Kerberoasting.

     We can utilise Linux for Kerberoasting with `impacket-GetUserSPNs`. We can run `sudo impacket-GetUserSPNs -request -dc-ip DC-IP fakecompany.local/USER` where 'USER' is a set of credentials we have found on the machine which once gives us the hash, which we can then pass to Hashcat will cracks it. It's worth noting that if this command fails due to "KRB_AP_ERR_SKEW(Clock skew too great)" we need to synchronise the time between our Kali machine and the DC which we can do with `rdate`.
     """
               kerberoasting += f"\n\nBy successfully Kerberoasting, you would have identified the accounts: `{vulns['vulns'][i]['SID1']}` and `{vulns['vulns'][i]['SID2']}`"

               doc.add_paragraph(kerberoasting)
               print("> Kerberoasting Added")
        elif (vulns['vulns'][i]['Name'] == "badACL"):
             doc.add_heading("BadACL's", 2)
             if (vulns['vulns'][i]['Status'] == "True"):         
                doc.add_paragraph("Bad ACLs were present on this machine which would have given the user an additional set of credentials to enumerate the network through.")
             else:
                  doc.add_paragraph("The ability to take advantage of BadACL vulnerabilities was not present on this machine.")
             print("> BadACL's Added")
        elif (vulns['vulns'][i]['Name'] == "ntlmRelay"):
             doc.add_heading("NTLM Relay", 2)
             if(vulns['vulns'][i]['Status'] == "True"):      
                  doc.add_paragraph("Using a combination of Responder and impacket-ntlmrelayx, it would be possible to gain a SMB shell as a random user through abusing this vulnerability.")
             else:
                  doc.add_paragraph("The ability to take advantage of an NTLM Relay was not present on this machine.")
             print("> NTLM Relay Added")        
    except IndexError:
        break

doc.save("results.docx")

print("> Report generation complete.")